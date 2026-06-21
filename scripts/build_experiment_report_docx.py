from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement as SharedOxmlElement
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "实验报告_截至2026年6月17日.docx"

LATIN_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
TITLE_FONT = "Microsoft YaHei"

COLOR_NAVY = RGBColor(31, 58, 95)
COLOR_BLUE = RGBColor(46, 116, 181)
COLOR_DARK = RGBColor(34, 34, 34)
COLOR_MUTED = RGBColor(96, 96, 96)
COLOR_BORDER = "C9D1DB"
COLOR_LIGHT_FILL = "F4F6F9"
COLOR_HEADER_FILL = "EAF0F7"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP = 80
CELL_MARGIN_BOTTOM = 80
CELL_MARGIN_START = 120
CELL_MARGIN_END = 120


def set_run_font(run, size: float | None = None, *, bold: bool | None = None, color: RGBColor | None = None,
                 italic: bool | None = None, font_name: str | None = None, east_asia_name: str | None = None) -> None:
    font_name = font_name or LATIN_FONT
    east_asia_name = east_asia_name or EAST_ASIA_FONT
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(paragraph, *, before: float = 0, after: float = 6, line_spacing: float = 1.15,
                         alignment: WD_ALIGN_PARAGRAPH | None = None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        paragraph.alignment = alignment


def set_style_fonts(style, *, size: float, bold: bool = False, color: RGBColor = COLOR_DARK,
                    font_name: str = LATIN_FONT, east_asia_name: str = EAST_ASIA_FONT,
                    before: float = 0, after: float = 6, line_spacing: float = 1.15) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))
    cell.width = Inches(width_dxa / 1440)


def set_table_borders(table, *, color: str = COLOR_BORDER, size: int = 8) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_table_cell_margins(table, *, top: int = CELL_MARGIN_TOP, bottom: int = CELL_MARGIN_BOTTOM,
                           start: int = CELL_MARGIN_START, end: int = CELL_MARGIN_END) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tbl_cell_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tbl_cell_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = TABLE_INDENT_DXA,
                       total_dxa: int = PAGE_WIDTH_DXA) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_table_cell_margins(table)
    set_table_borders(table)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "nil")


def add_text(paragraph, text: str, *, size: float = 11, bold: bool = False, color: RGBColor = COLOR_DARK,
             italic: bool = False, font_name: str = LATIN_FONT, east_asia_name: str = EAST_ASIA_FONT):
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size=size,
        bold=bold,
        color=color,
        italic=italic,
        font_name=font_name,
        east_asia_name=east_asia_name,
    )
    return run


def add_body_paragraph(doc: Document, text: str, *, before: float = 0, after: float = 6,
                       line_spacing: float = 1.15) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    set_paragraph_format(paragraph, before=before, after=after, line_spacing=line_spacing)
    add_text(paragraph, text)


def add_label_paragraph(doc: Document, label: str, text: str, *, after: float = 4) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    set_paragraph_format(paragraph, before=0, after=after, line_spacing=1.15)
    add_text(paragraph, f"{label}", bold=True, color=COLOR_NAVY)
    add_text(paragraph, text)


def add_section_title(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    add_text(paragraph, text, size={1: 16, 2: 13, 3: 12}[level], bold=True, color={1: COLOR_BLUE, 2: COLOR_BLUE, 3: COLOR_NAVY}[level])


def add_table_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=4, after=4, line_spacing=1.0)
    add_text(paragraph, text, size=10, color=COLOR_MUTED, italic=True)


def add_data_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers):
        set_cell_shading(cell, COLOR_HEADER_FILL)
        paragraph = cell.paragraphs[0]
        set_paragraph_format(paragraph, before=0, after=0, line_spacing=1.05)
        add_text(paragraph, header, size=10.5, bold=True, color=COLOR_NAVY)
    for row_data in rows:
        row_cells = table.add_row().cells
        for cell, value in zip(row_cells, row_data):
            paragraph = cell.paragraphs[0]
            set_paragraph_format(paragraph, before=0, after=0, line_spacing=1.08)
            add_text(paragraph, value, size=10.3)


def add_callout(doc: Document, title: str, body_lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PAGE_WIDTH_DXA], indent_dxa=0)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLOR_LIGHT_FILL)
    paragraph = cell.paragraphs[0]
    set_paragraph_format(paragraph, before=0, after=4, line_spacing=1.1)
    add_text(paragraph, title, size=11.5, bold=True, color=COLOR_NAVY)
    for line in body_lines:
        paragraph = cell.add_paragraph()
        set_paragraph_format(paragraph, before=0, after=2, line_spacing=1.1)
        add_text(paragraph, line, size=10.5)


def add_metadata_block(doc: Document) -> None:
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [1800, 7560], indent_dxa=0)
    remove_table_borders(table)
    entries = [
        ("报告主题", "WebAttackSim 两阶段实验报告（抽象训练阶段 + 真实 VulnHub 推理阶段）"),
        ("时间范围", "截至 2026 年 6 月 17 日已完成实验"),
        ("资料来源", "REPORT_STAGE1_RL_PRM_robustness.md、REPORT_STAGE2_vulnhub_inference.md 及对应 outputs/*.json"),
        ("报告目的", "把现有实验结果整理成适合汇报、归档和后续扩展的正式 Word 版本，并突出可复现结论与诚实局限。"),
    ]
    for row, (label, value) in zip(table.rows, entries):
        left, right = row.cells
        p_left = left.paragraphs[0]
        p_right = right.paragraphs[0]
        set_paragraph_format(p_left, before=0, after=0, line_spacing=1.0)
        set_paragraph_format(p_right, before=0, after=0, line_spacing=1.1)
        add_text(p_left, label, size=10.5, bold=True, color=COLOR_NAVY)
        add_text(p_right, value, size=10.5)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_fonts(normal, size=11, color=COLOR_DARK)

    heading1 = doc.styles["Heading 1"]
    set_style_fonts(heading1, size=16, bold=True, color=COLOR_BLUE, before=16, after=8, line_spacing=1.1,
                    font_name=LATIN_FONT, east_asia_name=TITLE_FONT)

    heading2 = doc.styles["Heading 2"]
    set_style_fonts(heading2, size=13, bold=True, color=COLOR_BLUE, before=12, after=6, line_spacing=1.08,
                    font_name=LATIN_FONT, east_asia_name=TITLE_FONT)

    heading3 = doc.styles["Heading 3"]
    set_style_fonts(heading3, size=12, bold=True, color=COLOR_NAVY, before=8, after=4, line_spacing=1.05,
                    font_name=LATIN_FONT, east_asia_name=TITLE_FONT)


def build_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    set_paragraph_format(title, before=0, after=6, line_spacing=1.0)
    add_text(title, "WebAttackSim 实验报告", size=23, bold=True, color=COLOR_DARK, east_asia_name=TITLE_FONT)

    subtitle = doc.add_paragraph()
    set_paragraph_format(subtitle, before=0, after=14, line_spacing=1.0)
    add_text(subtitle, "截至 2026 年 6 月 17 日的阶段一训练实验与阶段二真实靶场实验整理版", size=13.5, color=COLOR_MUTED)

    add_metadata_block(doc)

    add_callout(
        doc,
        "执行摘要",
        [
            "本项目已经形成“抽象训练 -> 真实适配 -> 真实靶场 A/B”两阶段实验链路，且关键结果能够互相闭环。",
            "阶段一证明：16 动作抽象环境能稳定产出可用的 Oracle 标签与 PRM 标签，Oracle 在诚实口径下具备目标导向排序能力，PRM 在排序、泛化和错误动作识别上表现稳定。",
            "阶段二证明：当前抽象动作 schema 对真实单主机 Web 攻击步骤的覆盖已足够高，真实瓶颈更多来自动作归一化器和提议器，而不是抽象空间本身。",
            "截至 2026 年 6 月 17 日，PRM 在真实靶场上呈现方向一致、幅度温和的正向收益，但样本量仍偏小，统计上尚不能下强结论。",
        ],
    )


def build_report(doc: Document) -> None:
    add_section_title(doc, "1. 研究目标与报告范围", 1)
    add_body_paragraph(
        doc,
        "本报告汇总 WebAttackSim 截至 2026 年 6 月 17 日已经完成的实验，目的是把分散在两份阶段性 Markdown 报告中的实验设置、关键结果和局限整理为一份可直接提交的 Word 正式版。报告只覆盖已经跑通且有数据支撑的内容，不把尚未完成的设想包装成结论。"
    )
    add_label_paragraph(doc, "核心问题一：", "抽象单主机 Web 攻击环境是否足够表达真实任务，并为过程奖励模型提供稳定、无泄漏的训练标签。")
    add_label_paragraph(doc, "核心问题二：", "在真实 VulnHub 靶场中，把第一阶段训练得到的 PRM 接到真实执行链路后，是否能带来稳定的重排序收益。")

    add_table_caption(doc, "表 1  本报告覆盖的实验阶段与评估重点")
    add_data_table(
        doc,
        ["阶段", "实验环境", "要回答的问题", "截至 2026-06-17 的结论"],
        [
            ["阶段一", "抽象单主机 Web 模拟环境", "Oracle 和 PRM 是否真的学到有用排序，而不是被掩码或规则抬高", "已回答；需用诚实指标口径解读"],
            ["阶段二", "真实 VulnHub 单主机 Web 靶场", "抽象动作是否覆盖真实步骤，PRM 是否有真实 uplift", "已跑通；有正向趋势，但统计功效不足"],
        ],
        [1100, 2200, 3200, 2860],
    )

    add_section_title(doc, "2. 总体方法与实验链路", 1)
    add_body_paragraph(
        doc,
        "整个项目采用两阶段设计。第一阶段在纯抽象环境中训练价值 Oracle 和 Pentest-PRM，用于回答“标签是否可靠”；第二阶段把冻结的一阶段产物接到真实靶场，用于回答“抽象是否能迁移，以及 PRM 是否在真实环境有帮助”。"
    )
    add_label_paragraph(doc, "阶段一链路：", "任务配置 -> 抽象状态转移 -> Web-RL 价值 Oracle -> PRM 标签 -> Pentest-PRM。")
    add_label_paragraph(doc, "阶段二链路：", "真实靶场输出 -> φ 状态解析 -> ψ 动作归一化 -> PRM 重排序 -> η 真正命令执行 -> φ 再观测。")
    add_callout(
        doc,
        "本报告坚持的口径",
        [
            "只把被数据直接支持的结果写成结论。",
            "凡是会被动作掩码、规则行、诱饵奖励或样本量不足抬高的数字，都明确降级说明，不作为主结论。",
            "所有“截至 6/17”的表述都在正文中写为“截至 2026 年 6 月 17 日”，避免日期歧义。",
        ],
    )

    add_section_title(doc, "3. 阶段一：抽象环境中的训练实验", 1)
    add_section_title(doc, "3.1 环境与任务集设计", 2)
    add_body_paragraph(
        doc,
        "阶段一只在抽象模拟器内进行，不涉及真实靶场。环境固定 16 类动作 schema，状态由 12 个可观测字段组成，目标文件、口令和 flag 等隐藏真值不会直接进入观测，只能通过环境反馈间接暴露，因此可以对后续“无泄漏”结论提供前提保障。"
    )
    add_table_caption(doc, "表 2  阶段一任务集与环境规模")
    add_data_table(
        doc,
        ["项目", "结果", "项目", "结果"],
        [
            ["动作 schema", "16 类固定动作", "任务总数", "65"],
            ["结构族数量", "12", "拓扑签名", "12"],
            ["难度分布", "easy 10 / medium 15 / hard 40", "链深范围", "2 到 12 步"],
            ["训练/留出切分", "45 训练 / 10 留出实例 / 10 留出链型", "留出链型", "leak_login、rce_privesc"],
        ],
        [2000, 2680, 2000, 2680],
    )
    add_body_paragraph(
        doc,
        "任务覆盖配置泄漏、弱口令、SQLi/LFI、RCE、文件上传与提权等典型 Web 攻击链，并且每个任务都带有专家计划、专家轨迹和诱饵路径，从而让“排序是否合理”成为可量化问题。"
    )

    add_section_title(doc, "3.2 Oracle 的主要结果", 2)
    add_body_paragraph(
        doc,
        "Oracle 使用 DQN 训练，其目标不是直接生成最终策略，而是为后续 PRM 提供更细粒度的动作价值标签。为了避免被掩码饱和误导，阶段一明确把“相对随机策略的增量”而不是“掩码下原始成功率”当作主结论。"
    )
    add_table_caption(doc, "表 3  Oracle 的关键指标（诚实口径）")
    add_data_table(
        doc,
        ["指标", "结果", "解释"],
        [
            ["专家动作 top-1 相对随机增量", "+0.127", "这是阶段一对 Oracle 最重要的 headline 指标。"],
            ["无掩码 permissive goal rate", "0.40 vs 随机 0.0125", "说明模型确实学到了任务相关能力，而不是只会跟着掩码走。"],
            ["goal-aligned Q* top-1 / top-3", "0.71 / 0.98", "去掉诱饵奖励伪影后，Oracle 对“朝目标推进”排序是合理的。"],
            ["goal-aligned Q* Spearman", "+0.32", "与 literal-reward Q* 的负相关不同，这个值更能反映真实目标导向。"],
            ["多种子排序一致率", "0.972", "标签来源在不同随机种子下保持高稳定性。"],
            ["hard mode 关键决策点相关", "-0.005 -> +0.373", "收紧预算后，同状态下动作选择开始真正影响结果，Oracle 的价值判断才变得有意义。"],
        ],
        [2200, 1700, 5460],
    )
    add_callout(
        doc,
        "为什么不把 masked goal rate 当主结论",
        [
            "掩码会先帮系统筛掉大量不合法动作，因此掩码内的原始成功率天然偏高。",
            "阶段一已经明确发现 masked goal rate 会出现饱和现象，因此本报告改用“相对随机增量”和“无掩码成功率”来表达 Oracle 是否真的学到东西。",
        ],
    )

    add_section_title(doc, "3.3 PRM 的主要结果", 2)
    add_body_paragraph(
        doc,
        "Pentest-PRM 使用结构化“状态 + 动作”特征训练，严格不把 Oracle 的 q 值作为输入特征，从源头上避免标签泄漏。其角色是逐步重排序器，而不是独立自治策略。"
    )
    add_table_caption(doc, "表 4  PRM 的关键指标")
    add_data_table(
        doc,
        ["指标", "结果", "解读"],
        [
            ["oracle 标注子集 pairwise", "0.890", "PRM 在可比候选动作之间具有较强排序能力。"],
            ["pairwise 95% bootstrap CI", "[0.843, 0.937]", "整体结论有置信区间支撑，而不是单个点估计。"],
            ["留出实例 pairwise", "0.980", "对同抽象空间中的新实例泛化很强。"],
            ["留出链型 pairwise", "0.800", "对新链型仍有效，但比同分布实例明显更难。"],
            ["零样本新任务族 pairwise", "0.873", "具备跨族迁移能力，但仍低于同分布参考 0.890。"],
            ["错误动作识别 ROC-AUC", "0.893", "更适合作为逐步筛选器，尤其适合压低明显错误动作。"],
        ],
        [2200, 1700, 5460],
    )
    add_body_paragraph(
        doc,
        "校准方面，sigmoid/Platt 后校准能把最难的留出链型 ECE 从 0.155 压到 0.067，但会损害原本已经校准较好的留出实例，因此结论不是“全局必须校准”，而是“只在最难切片选择性使用后校准”。"
    )
    add_body_paragraph(
        doc,
        "阶段一的另一项重要发现是：PRM 接入闭环自治 rollout 并不会自动带来策略层面的收益。原因不是 PRM 无用，而是它本质上只负责候选动作之间的相对排序；如果上游 proposer 根本没有提出正确 exploit，PRM 也无法凭空生成正确动作。"
    )

    add_section_title(doc, "3.4 阶段一结论", 2)
    add_label_paragraph(doc, "结论一：", "抽象环境设计已经具备足够的任务覆盖和可解释性，能够支撑阶段一的训练与评测。")
    add_label_paragraph(doc, "结论二：", "Oracle 虽然仍然偏弱、且对掩码敏感，但在诚实口径下已经具备稳定的目标导向排序能力。")
    add_label_paragraph(doc, "结论三：", "PRM 在排序质量、泛化能力和错误动作识别上都表现稳健，但它应被定位为“重排序器”，而不是独立策略模型。")

    doc.add_page_break()

    add_section_title(doc, "4. 阶段二：真实 VulnHub 靶场实验", 1)
    add_section_title(doc, "4.1 真实靶场的适配与安全边界", 2)
    add_body_paragraph(
        doc,
        "阶段二的目标不是再证明阶段一有没有学到，而是测试这些抽象训练产物能否迁移到真实单主机 Web 靶场。所有实验都只在自有、隔离、绑定到 127.0.0.1 的训练容器中运行，并且通过 `AuthorizationGate`、命令白名单和审计日志约束执行范围。真实执行只允许 `id`、`whoami`、`cat /etc/passwd` 等只读命令。"
    )
    add_label_paragraph(doc, "φ 的作用：", "把真实工具输出解析回抽象状态。")
    add_label_paragraph(doc, "ψ 的作用：", "把 LLM 的动作表述归一化为 16 类抽象动作之一。")
    add_label_paragraph(doc, "η 的作用：", "把抽象动作转换成真实可执行的、受安全闸约束的命令。")

    add_section_title(doc, "4.2 离线抽象差距评估", 2)
    add_body_paragraph(
        doc,
        "在 7 个手工标注的 VulnHub 级 fixture、共 71 个步骤上，阶段二先离线评估“抽象是否够用”，避免直接把 sim-to-real 问题全部堆给在线实验。结果显示，真正的瓶颈并不是 16 动作 schema 本身，而是动作归一化器 ψ。"
    )
    add_table_caption(doc, "表 5  阶段二离线抽象差距评估")
    add_data_table(
        doc,
        ["指标", "结果", "解释"],
        [
            ["out-of-abstraction", "8.5%", "16 动作抽象大约覆盖了 92% 的真实步骤，因此不需要立刻扩 schema。"],
            ["φ 状态还原召回", "94.8%", "状态解析器已经足够可靠，可以把真实反馈重新映射回抽象状态。"],
            ["ψ 准确率", "49% -> 78.5%", "经过阶段二增强后，真实措辞映射显著改善，瓶颈从“抽象空间”转向“动作归一化器”。"],
            ["越界步骤类型", "SSH 登录、离线破哈希、SMB/二进制利用、su 切换", "这些缺口基本都属于“非 Web 原语”，本就在单主机 Web 范围之外。"],
        ],
        [2200, 1700, 5460],
    )

    add_section_title(doc, "4.3 真实靶场与端到端打通结果", 2)
    add_body_paragraph(
        doc,
        "截至 2026 年 6 月 17 日，阶段二已经在 4 个真实 Vulhub 单主机 Web 容器上完成端到端链路验证。固定序列的 live smoke 说明：η 可以生成靶机特定命令，受闸执行器可以安全执行，φ 能够把真实响应重新解释成抽象状态。"
    )
    add_table_caption(doc, "表 6  真实靶场配置与端到端结果")
    add_data_table(
        doc,
        ["靶机", "漏洞类型", "是否有明显框架标识", "端到端结果"],
        [
            ["ThinkPHP 5-rce", "ThinkPHP 5.0.x 框架 RCE", "是", "可达命令执行，身份为 www-data"],
            ["ThinkPHP 5.0.23", "ThinkPHP 5.0.23 框架 RCE", "是", "可达命令执行，身份为 www-data"],
            ["Struts2 S2-048", "OGNL RCE（CVE-2017-9791）", "是", "可达 root 命令执行，并读取 /etc/passwd"],
            ["php-cgi 2012-1823", "PHP-CGI 参数注入 RCE", "否", "固定序列可打通，但自治发现失败"],
        ],
        [2200, 2900, 1200, 3060],
    )
    add_body_paragraph(
        doc,
        "真实靶场还暴露出多类 sim-to-real 工程问题，例如 Windows 默认 GBK 解码导致真实响应崩溃、HTML/CSS 文本被误识别为凭据、curl 对 `[]` 的 globbing 等。这些问题在阶段二都已经被定位并修复，也说明阶段二的价值不仅是“跑分”，更是暴露真实适配缺口。"
    )

    add_section_title(doc, "4.4 自治 A/B 实验结果", 2)
    add_body_paragraph(
        doc,
        "真实 A/B 实验使用 `deepseek-v4-pro` 作为动作提议器，每个靶机、每个实验臂运行 6 次，预算 14，目标定义为“获得命令执行并读取敏感文件”。A 组由 PRM 对候选动作重排序后选首位，B 组直接采用 proposer 原始顺序。"
    )
    add_table_caption(doc, "表 7  真实靶场 A/B 结果（每靶每臂 6 次）")
    add_data_table(
        doc,
        ["靶机", "PRM 组", "基线组", "差值", "结果解读"],
        [
            ["ThinkPHP 5-rce", "6/6 = 100%", "3/6 = 50%", "+50pp", "PRM 明显提高达成目标的可靠性。"],
            ["ThinkPHP 5.0.23", "4/6 = 67%", "4/6 = 67%", "0", "说明单靶高分并不总能复制。"],
            ["Struts2 S2-048", "5/6 = 83%", "4/6 = 67%", "+16pp", "仍有正向收益，但幅度更温和。"],
            ["php-cgi 2012-1823", "0/6 = 0%", "0/6 = 0%", "0", "问题不在 PRM，而在 proposer 没有提出 exploit。"],
        ],
        [2200, 1500, 1500, 900, 3260],
    )
    add_table_caption(doc, "表 8  聚合结果与统计判读")
    add_data_table(
        doc,
        ["统计范围", "PRM 成功率", "基线成功率", "结论"],
        [
            ["3 个“自报家门”靶", "15/18 = 83.3%", "11/18 = 61.1%", "提升 +22.2pp，但 p = 0.137，方向正确而统计上尚不显著。"],
            ["全部 4 个靶", "15/24 = 62.5%", "11/24 = 45.8%", "提升 +16.7pp，p = 0.247，说明当前样本量仍不足。"],
        ],
        [2100, 1800, 1800, 3660],
    )
    add_callout(
        doc,
        "阶段二最重要的真实结论",
        [
            "PRM 在每一个真实靶上都没有比基线更差，说明重排序方向基本可靠。",
            "当前 uplift 主要体现在“提高任务完成的稳定性”，而不是单次速度更快。",
            "当 proposer 根本没有提出正确 exploit 时，PRM 作为 reranker 没有能力凭空补出正确动作，这也是 php-cgi 靶机双臂都为 0 的原因。",
        ],
    )

    add_section_title(doc, "5. 综合结论与局限", 1)
    add_body_paragraph(
        doc,
        "把两阶段实验连起来看，截至 2026 年 6 月 17 日，WebAttackSim 已经完成了从抽象训练到真实单主机 Web 靶场的最小闭环验证。阶段一回答了“标签是否可靠”，阶段二回答了“抽象是否能迁移，以及 PRM 是否在真实靶上有帮助”。"
    )
    add_table_caption(doc, "表 9  截至 2026-06-17 的综合结论")
    add_data_table(
        doc,
        ["问题", "当前结论", "对后续工作的含义"],
        [
            ["抽象动作是否够用", "基本够用", "8.5% 的越界率说明短期不必扩 schema，优先继续打磨 ψ 与 proposer。"],
            ["Oracle 标签是否可用", "可用，但不能夸大", "需要继续坚持诚实指标，不把掩码下高成功率当 headline。"],
            ["PRM 是否能迁移到真实靶场", "能迁移，并有温和正向收益", "后续重点是扩样本量、提高显著性，而不是推翻方法。"],
            ["当前最大瓶颈是什么", "proposer 与真实世界可识别性", "PRM 是 reranker，不是 proposer；无 banner 漏洞需要主动探测能力。"],
        ],
        [2100, 2900, 4360],
    )
    add_table_caption(doc, "表 10  需要如实保留的局限")
    add_data_table(
        doc,
        ["局限点", "截至 2026-06-17 的具体表现", "建议动作"],
        [
            ["样本量不足", "每靶每臂只有 6 次，真实 uplift 尚未达到显著", "扩到每靶约 20 次，并增加更多框架类型靶机。"],
            ["proposer 先于 PRM 成为瓶颈", "php-cgi 无 banner 时，双臂都为 0", "补主动探测动作与漏洞发现能力，而不是只优化 PRM。"],
            ["PRM 不是独立策略", "闭环自治 rollout 不体现其真实价值", "继续把 PRM 定位为候选动作排序器。"],
            ["阶段二增强 ψ 仍未回流到 PRM 训练", "若把增强版 ψ 纳入特征路径，需要重训 PRM", "后续单独安排一次“增强 ψ + PRM 重训”实验。"],
        ],
        [1700, 4260, 3400],
    )

    add_section_title(doc, "6. 复现与材料索引", 1)
    add_body_paragraph(
        doc,
        "为了便于后续继续更新到 2026 年 6 月 17 日之后的新实验，下面保留本报告直接使用的关键材料和命令。"
    )
    add_label_paragraph(doc, "材料 1：", "REPORT_STAGE1_RL_PRM_robustness.md，用于追溯阶段一的实验设定、结果和局限。")
    add_label_paragraph(doc, "材料 2：", "REPORT_STAGE2_vulnhub_inference.md，用于追溯阶段二真实靶场实验的完整描述。")
    add_label_paragraph(doc, "命令 1：", "python scripts/run_training_stage.py，用于复现阶段一主链路；需要更完整检查时再加 --include-slow。")
    add_label_paragraph(doc, "命令 2：", "python -m stage2.preflight，用于检查阶段二环境与安全前置条件。")
    add_label_paragraph(doc, "命令 3：", "python -m stage2.aggregate_multibox，用于输出多靶 A/B 聚合结果。")
    add_body_paragraph(
        doc,
        "如果后续要继续补 2026 年 6 月 17 日之后的新实验，建议沿用本报告结构追加更新：先补阶段二统计功效，再补 proposer/主动探测能力，最后再考虑扩动作 schema。"
    )


def write_document(output_path: Path) -> Path:
    doc = Document()
    configure_document(doc)
    build_cover(doc)
    doc.add_page_break()
    build_report(doc)
    doc.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the consolidated WebAttackSim experiment report as a DOCX.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="Output DOCX path.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output_path = args.output.resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    write_document(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
